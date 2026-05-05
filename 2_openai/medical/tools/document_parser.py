"""
tools/document_parser.py — Document Parsing Tool for MediScan AI

ROOT CAUSE OF THE TATA 1MG FORMATTING PROBLEM:
    PyMuPDF get_text("text") extracts each PDF table cell on its own line:
        Hemoglobin        <- test name alone
        15.3              <- value alone
        g/dL              <- unit alone
        13.0-17.0         <- reference range alone

    THE FIX — Block-based extraction with Y-coordinate row grouping:
    get_text("blocks") returns each text block with bounding box (x0,y0,x1,y1).
    Blocks sharing the same Y position are in the same table row.
    We group by Y proximity and join cells with " | ":
        Hemoglobin | 15.3 | g/dL | 13.0-17.0 | Cyanide Free SLS
"""

import re
import os
from dataclasses import dataclass, field

import fitz
from docx import Document as DocxDocument

from utils.validator import validate_file, ValidationResult
from utils.sanitizer import sanitize, get_text_stats


# ─────────────────────────────────────────────────────────────
#  ParsedDocument — structured output of this tool
# ─────────────────────────────────────────────────────────────

@dataclass
class ParsedDocument:
    success:       bool
    text:          str
    raw_text:      str   = ""
    file_name:     str   = ""
    extension:     str   = ""
    page_count:    int   = 0
    word_count:    int   = 0
    char_count:    int   = 0
    chunk_count:   int   = 1
    is_meaningful: bool  = False
    warning:       str   = ""
    error:         str   = ""


# ─────────────────────────────────────────────────────────────
#  _extract_page_text()  — block-based PDF page extraction
#
#  WHY BLOCKS INSTEAD OF get_text("text")?
#
#  get_text("text") reads PDF content top-to-bottom in glyph order.
#  For multi-column tables this means each cell becomes its own line.
#
#  get_text("blocks") returns (x0, y0, x1, y1, text, block_no, type).
#  Blocks with the same Y position = same table row.
#  We sort by Y then X, group blocks within Y_TOLERANCE pixels,
#  and join each row's cells with " | ".
#
#  Result: "Hemoglobin | 15.3 | g/dL | 13.0-17.0 | Cyanide Free SLS"
#  The LLM can now correctly parse each complete lab parameter row.
# ─────────────────────────────────────────────────────────────

Y_TOLERANCE = 5  # pixels — blocks within this Y distance = same row


def _extract_page_text(page) -> str:
    """
    Extract text from a PDF page using block-based row reconstruction.
    Groups text blocks by vertical position to reconstruct table rows.
    """
    blocks = page.get_text("blocks")
    # Filter: text blocks only (type==0), non-empty
    text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]

    if not text_blocks:
        return ""

    # Sort top-to-bottom, then left-to-right
    text_blocks.sort(key=lambda b: (b[1], b[0]))

    # Group into rows by Y proximity
    rows: list[list] = []
    current_row: list = [text_blocks[0]]
    current_y: float = text_blocks[0][1]

    for block in text_blocks[1:]:
        if abs(block[1] - current_y) <= Y_TOLERANCE:
            current_row.append(block)
        else:
            rows.append(current_row)
            current_row = [block]
            current_y = block[1]
    if current_row:
        rows.append(current_row)

    # Build output lines
    lines: list[str] = []
    for row in rows:
        # Sort cells left-to-right
        row.sort(key=lambda b: b[0])
        cells = []
        for block in row:
            # Normalize internal newlines within a cell to spaces
            cell = re.sub(r'\s*\n\s*', ' ', block[4]).strip()
            if cell:
                cells.append(cell)

        if not cells:
            continue
        if len(cells) == 1:
            lines.append(cells[0])
        else:
            lines.append(" | ".join(cells))

    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────
#  _parse_pdf()
# ─────────────────────────────────────────────────────────────

def _parse_pdf(file_path: str, file_name: str) -> ParsedDocument:
    doc = None
    try:
        doc = fitz.open(file_path)
        page_count = len(doc)
        page_texts = []

        for page_num, page in enumerate(doc, start=1):
            page_text = _extract_page_text(page)
            if page_text.strip():
                page_texts.append(f"[Page {page_num}]\n{page_text}")

        raw_text = "\n\n".join(page_texts)
        clean_text = sanitize(raw_text)
        stats = get_text_stats(clean_text)

        warning = ""
        if not stats["is_meaningful"] and page_count > 0:
            warning = (
                "⚠️ This PDF appears to be a scanned image. "
                "Text extraction returned minimal content. "
                "RC2 will support scanned documents via OCR."
            )

        return ParsedDocument(
            success=True,
            raw_text=raw_text,
            text=clean_text,
            file_name=file_name,
            extension="pdf",
            page_count=page_count,
            word_count=stats["word_count"],
            char_count=stats["char_count"],
            chunk_count=stats["chunk_count"],
            is_meaningful=stats["is_meaningful"],
            warning=warning,
        )

    except fitz.FileDataError as e:
        return ParsedDocument(
            success=False, text="", file_name=file_name, extension="pdf",
            error=f"PDF appears to be corrupted or is not a valid PDF file. ({e})",
        )
    except Exception as e:
        return ParsedDocument(
            success=False, text="", file_name=file_name, extension="pdf",
            error=f"Unexpected error reading PDF: ({e})",
        )
    finally:
        if doc is not None:
            doc.close()


# ─────────────────────────────────────────────────────────────
#  _parse_docx()
# ─────────────────────────────────────────────────────────────

def _parse_docx(file_path: str, file_name: str, extension: str) -> ParsedDocument:
    try:
        doc = DocxDocument(file_path)
        collected_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                collected_parts.append(text)

        for table_idx, table in enumerate(doc.tables, start=1):
            collected_parts.append(f"\n[Table {table_idx}]")
            for row in table.rows:
                cell_texts = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if cell_texts:
                    collected_parts.append(" | ".join(cell_texts))

        raw_text = "\n".join(collected_parts)
        clean_text = sanitize(raw_text)
        stats = get_text_stats(clean_text)
        estimated_pages = max(1, stats["word_count"] // 300)

        warning = ""
        if not stats["is_meaningful"]:
            warning = (
                "⚠️ This document contains very little text. "
                "It may be empty or contain only images/embedded objects."
            )

        return ParsedDocument(
            success=True,
            text=clean_text,
            raw_text=raw_text,
            file_name=file_name,
            extension=extension,
            page_count=estimated_pages,
            word_count=stats["word_count"],
            char_count=stats["char_count"],
            chunk_count=stats["chunk_count"],
            is_meaningful=stats["is_meaningful"],
            warning=warning,
        )

    except Exception as e:
        error_message = (
            "Legacy .doc format could not be read. Please resave as .docx."
            if extension == "doc" else str(e)
        )
        return ParsedDocument(
            success=False, text="", file_name=file_name, extension=extension,
            error=f"Error reading Word document: {error_message}",
        )


# ─────────────────────────────────────────────────────────────
#  parse_document()  — MAIN PUBLIC FUNCTION
# ─────────────────────────────────────────────────────────────

def parse_document(file: object) -> ParsedDocument:
    """
    Main entry point: validate and parse an uploaded PDF or DOCX file.
    Always returns a ParsedDocument — never raises exceptions.
    """
    validation: ValidationResult = validate_file(file)

    if not validation.is_valid:
        return ParsedDocument(
            success=False,
            text="",
            file_name=validation.file_name,
            extension=validation.extension,
            error=validation.message,
        )

    file_path = validation.metadata.get("file_path", "")
    file_name = validation.file_name
    extension = validation.extension

    if extension == "pdf":
        return _parse_pdf(file_path, file_name)
    elif extension in ("docx", "doc"):
        return _parse_docx(file_path, file_name, extension)
    else:
        return ParsedDocument(
            success=False, text="", file_name=file_name, extension=extension,
            error=f"Internal error: unhandled file type '.{extension}'.",
        )


# ─────────────────────────────────────────────────────────────
#  format_parsed_for_display()  — UI helper
# ─────────────────────────────────────────────────────────────

def format_parsed_for_display(parsed: ParsedDocument) -> str:
    """Format a ParsedDocument into markdown for the Gradio Raw Text tab."""
    if not parsed.success:
        return f"## ❌ Parsing Failed\n\n**Error:** {parsed.error}"

    lines = [
        "## 📄 Extracted Raw Text", "",
        "| Property | Value |", "|---|---|",
        f"| **File** | {parsed.file_name} |",
        f"| **Format** | .{parsed.extension.upper()} |",
        f"| **Pages** | {parsed.page_count} |",
        f"| **Words extracted** | {parsed.word_count:,} |",
        f"| **Characters** | {parsed.char_count:,} |",
        f"| **LLM chunks** | {parsed.chunk_count} |",
        "",
    ]

    if parsed.warning:
        lines += [f"> {parsed.warning}", ""]

    lines += [
        "### Raw Content", "```",
        parsed.raw_text[:3000] + ("..." if len(parsed.raw_text) > 3000 else ""),
        "```",
    ]

    return "\n".join(lines)
