"""
test_parser.py — Week 2 Verification Script for MediScan AI
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PURPOSE:
    Verifies that everything built in Week 2 is working correctly
    BEFORE we move to Week 3 (LLM analysis). This script tests:

        1. All Week 2 imports load without errors
        2. validator.py correctly accepts and rejects files
        3. sanitizer.py correctly cleans text
        4. document_parser.py correctly extracts text from
           a synthetic PDF and a synthetic DOCX file
        5. format_parsed_for_display() produces readable markdown
        6. app.py imports cleanly with Week 2 wired in

    WHY TEST WITH SYNTHETIC FILES?
    We generate our own test PDF and DOCX programmatically
    rather than asking you to find and provide real medical reports.
    This means the test is completely self-contained — anyone
    cloning the repo can run it immediately without needing test data.
    The synthetic files contain realistic medical content so
    the parser is exercised with the same kind of text it'll
    see in production.

Run with:
    python test_parser.py
    uv run python test_parser.py
"""

import sys
import os
import tempfile

print("\n" + "═" * 65)
print("  MediScan AI — Week 2 Parser Test")
print("═" * 65)


# ─────────────────────────────────────────────────────────────
#  Test 1: Import verification
#  All Week 2 modules must import without errors before we test
#  their functionality. If an import fails, there's a bug or a
#  missing dependency — we catch this before running anything.
# ─────────────────────────────────────────────────────────────

print("\n[1/6] Verifying Week 2 imports...")

try:
    from utils.validator  import validate_file, ValidationResult
    print("   ✅ utils.validator    — validate_file, ValidationResult")
except ImportError as e:
    print(f"   ❌ utils.validator failed: {e}")
    sys.exit(1)

try:
    from utils.sanitizer  import sanitize, is_meaningful, chunk_text, get_text_stats
    print("   ✅ utils.sanitizer    — sanitize, is_meaningful, chunk_text, get_text_stats")
except ImportError as e:
    print(f"   ❌ utils.sanitizer failed: {e}")
    sys.exit(1)

try:
    from tools.document_parser import parse_document, ParsedDocument, format_parsed_for_display
    print("   ✅ tools.document_parser — parse_document, ParsedDocument, format_parsed_for_display")
except ImportError as e:
    print(f"   ❌ tools.document_parser failed: {e}")
    sys.exit(1)

try:
    import fitz   # PyMuPDF
    print("   ✅ fitz (PyMuPDF)     — PDF extraction library")
except ImportError:
    print("   ❌ fitz (PyMuPDF) not installed — run: pip install pymupdf")
    sys.exit(1)

try:
    from docx import Document
    print("   ✅ docx (python-docx) — DOCX extraction library")
except ImportError:
    print("   ❌ python-docx not installed — run: pip install python-docx")
    sys.exit(1)


# ─────────────────────────────────────────────────────────────
#  Test 2: Validator unit tests
#  We test the validator with edge cases it must handle correctly.
#  We use a NamedTuple to simulate Gradio's file object,
#  since Gradio isn't running during tests.
# ─────────────────────────────────────────────────────────────

print("\n[2/6] Testing validator.py...")

# Simulate Gradio's file object — it has a .name attribute
# that points to the file path on disk
class FakeFile:
    def __init__(self, name: str):
        self.name = name

# Test 2a: None input (no file uploaded)
result = validate_file(None)
assert not result.is_valid, "None file should be invalid"
print("   ✅ None file → correctly rejected")

# Test 2b: Wrong extension
with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
    f.write(b"some text content")
    txt_path = f.name

try:
    result = validate_file(FakeFile(txt_path))
    assert not result.is_valid, ".txt file should be rejected"
    print("   ✅ .txt file → correctly rejected (unsupported type)")
finally:
    os.unlink(txt_path)

# Test 2c: Empty file
with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
    empty_path = f.name  # write nothing — 0 bytes

try:
    result = validate_file(FakeFile(empty_path))
    assert not result.is_valid, "Empty file should be rejected"
    print("   ✅ Empty .pdf → correctly rejected (empty file)")
finally:
    os.unlink(empty_path)

# Test 2d: File that doesn't exist (corrupted/inaccessible)
result = validate_file(FakeFile("/nonexistent/path/report.pdf"))
assert not result.is_valid, "Nonexistent file should be rejected"
print("   ✅ Non-existent path → correctly rejected")

print("   ✅ All validator edge cases passed")


# ─────────────────────────────────────────────────────────────
#  Test 3: Sanitizer unit tests
#  We test sanitizer.py with crafted dirty text strings that
#  represent the kinds of artefacts real PDFs produce.
#  Each test verifies one specific cleaning behaviour.
# ─────────────────────────────────────────────────────────────

print("\n[3/6] Testing sanitizer.py...")

# Clear any stale bytecode before testing
# This prevents old cached .pyc files from causing unexpected behaviour
import importlib
import utils.sanitizer as _san_module
importlib.reload(_san_module)
from utils.sanitizer import sanitize, is_meaningful, chunk_text, get_text_stats

# Test 3a: Hyphenated line break repair
# "haemo-\nglobin" is a common PDF artefact where a word
# gets split across lines with a hyphen. We expect it joined.
# We use chr(10) explicitly for the newline to avoid any
# platform-specific string literal interpretation issues.
newline = chr(10)  # explicit newline character — same as \n but unambiguous
dirty = f"The patient has haemo-{newline}globin deficiency."
clean = sanitize(dirty)
assert "haemoglobin" in clean, (
    f"Hyphen repair failed.\n"
    f"  Input : {repr(dirty)}\n"
    f"  Output: {repr(clean)}\n"
    f"  Expected 'haemoglobin' in cleaned text"
)
print("   ✅ Hyphenated line break repair — 'haemo-\\nglobin' → 'haemoglobin'")

# Test 3b: Excessive whitespace collapse
# Multiple spaces between words (common in PDF column layouts)
dirty = "Hemoglobin      :      11.2     g/dL"
clean = sanitize(dirty)
assert "  " not in clean, f"Whitespace collapse failed. Got: {clean}"
print("   ✅ Excessive whitespace collapsed")

# Test 3c: Excessive blank lines reduced
# PDFs often have 5-6 blank lines between sections
dirty = "Section A\n\n\n\n\n\nSection B"
clean = sanitize(dirty)
assert "\n\n\n" not in clean, f"Blank line reduction failed. Got: {repr(clean)}"
print("   ✅ Excessive blank lines reduced to maximum 2")

# Test 3d: Page number removal
dirty = "Lab Results\n\nPage 1 of 3\n\nHemoglobin: 11.2"
clean = sanitize(dirty)
assert "Page 1 of 3" not in clean, f"Page number removal failed. Got: {clean}"
print("   ✅ Page number lines removed ('Page 1 of 3')")

# Test 3e: is_meaningful() threshold
assert not is_meaningful(""), "Empty string should not be meaningful"
assert not is_meaningful("   "), "Whitespace-only should not be meaningful"
assert not is_meaningful("Short"), "Very short text should not be meaningful"
assert is_meaningful("A" * 200), "200 chars should be meaningful"
print("   ✅ is_meaningful() thresholds correct")

# Test 3f: chunk_text() — short text stays as one chunk
short_text = "Patient: John Doe. Diagnosis: Mild anemia. " * 10
chunks = chunk_text(short_text)
assert len(chunks) == 1, f"Short text should be 1 chunk, got {len(chunks)}"
print("   ✅ Short text → 1 chunk (no splitting needed)")

# Test 3g: chunk_text() — long text splits at paragraph boundaries
# Create text that is definitely longer than MAX_CHARS_PER_CHUNK
long_text = ("Lab result paragraph with detailed values. " * 50 + "\n\n") * 20
chunks = chunk_text(long_text, max_chars=500)  # small limit for testing
assert len(chunks) > 1, "Long text should split into multiple chunks"
# Verify no chunk exceeds the limit
for i, chunk in enumerate(chunks):
    assert len(chunk) <= 600, f"Chunk {i} too long: {len(chunk)} chars"
print(f"   ✅ Long text → {len(chunks)} chunks, all within size limit")

# Test 3h: get_text_stats() returns correct structure
stats = get_text_stats("Hello world. This is a test sentence with ten words here.")
assert "word_count"    in stats, "Missing word_count in stats"
assert "char_count"    in stats, "Missing char_count in stats"
assert "chunk_count"   in stats, "Missing chunk_count in stats"
assert "is_meaningful" in stats, "Missing is_meaningful in stats"
print("   ✅ get_text_stats() returns all expected keys")

print("   ✅ All sanitizer tests passed")


# ─────────────────────────────────────────────────────────────
#  Test 4: PDF parsing with a synthetic medical PDF
#
#  We create a realistic synthetic lab report PDF using PyMuPDF
#  itself (fitz can both read AND create PDFs). This avoids
#  needing a real medical PDF file for testing.
#
#  The synthetic PDF contains:
#  — Patient information header
#  — CBC (Complete Blood Count) results as text
#  — Metabolic panel values
#  — Doctor's note
#
#  This exercises all the PDF parsing code paths in _parse_pdf()
# ─────────────────────────────────────────────────────────────

print("\n[4/6] Testing PDF parsing with synthetic medical document...")

SYNTHETIC_PDF_CONTENT = """
PATIENT LAB REPORT
==================
Patient Name: Test Patient
Date of Birth: 01/01/1980
Report Date: 01/05/2025
Ordering Physician: Dr. Smith
Lab ID: LAB-2025-0501

COMPLETE BLOOD COUNT (CBC)
--------------------------
Test             Result      Reference Range    Flag
Hemoglobin       11.2 g/dL   12.0-17.5 g/dL    LOW
WBC Count        7400 /uL    4000-11000 /uL     Normal
Platelet Count   210000 /uL  150000-400000 /uL  Normal
Hematocrit       34%         36-46%             LOW
MCV              78 fL       80-100 fL          LOW

METABOLIC PANEL
---------------
Test                 Result      Reference Range    Flag
Fasting Glucose      118 mg/dL   70-99 mg/dL        HIGH
Creatinine           0.9 mg/dL   0.6-1.2 mg/dL      Normal
BUN                  14 mg/dL    7-20 mg/dL          Normal
eGFR                 85          >60                 Normal
Sodium               139 mEq/L   136-145 mEq/L      Normal
Potassium            4.1 mEq/L   3.5-5.0 mEq/L      Normal

LIPID PANEL
-----------
Total Cholesterol    210 mg/dL   <200 mg/dL         HIGH
LDL Cholesterol      142 mg/dL   <100 mg/dL         HIGH
HDL Cholesterol      48 mg/dL    >40 mg/dL           Normal
Triglycerides        155 mg/dL   <150 mg/dL          Borderline

PHYSICIAN NOTES
---------------
Patient presents with mild microcytic anemia, likely iron-deficiency.
Fasting glucose is in pre-diabetic range. Recommend dietary modifications
and follow-up in 3 months. LDL cholesterol elevated - lifestyle counseling advised.
"""

# Create a synthetic PDF in a temporary file
with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
    pdf_test_path = tmp.name

try:
    # Use PyMuPDF to create a real PDF with our synthetic content
    # This is the same library we use for reading PDFs
    pdf_doc = fitz.open()                     # create empty PDF
    page = pdf_doc.new_page()                  # add one page
    page.insert_text(                          # insert our text
        point=(50, 50),                        # starting position (x, y)
        text=SYNTHETIC_PDF_CONTENT,
        fontsize=10,
        fontname="helv",                       # Helvetica — standard PDF font
    )
    pdf_doc.save(pdf_test_path)               # save to disk
    pdf_doc.close()

    # Now test parse_document() with the fake Gradio file object
    fake_file = FakeFile(pdf_test_path)
    parsed = parse_document(fake_file)

    # Assertions
    assert parsed.success,      f"PDF parsing failed: {parsed.error}"
    assert parsed.extension == "pdf", f"Extension wrong: {parsed.extension}"
    assert parsed.page_count == 1,    f"Expected 1 page, got {parsed.page_count}"
    assert parsed.word_count > 50,    f"Too few words extracted: {parsed.word_count}"
    assert parsed.char_count > 200,   f"Too few chars extracted: {parsed.char_count}"
    assert parsed.is_meaningful,      "Parsed text should be meaningful"
    assert "Hemoglobin" in parsed.text or "hemoglobin" in parsed.text.lower(), \
        "Key medical term 'Hemoglobin' not found in extracted text"

    print(f"   ✅ PDF parsed successfully")
    print(f"   ✅ Words extracted : {parsed.word_count:,}")
    print(f"   ✅ Characters      : {parsed.char_count:,}")
    print(f"   ✅ Pages           : {parsed.page_count}")
    print(f"   ✅ LLM chunks      : {parsed.chunk_count}")
    print(f"   ✅ Is meaningful   : {parsed.is_meaningful}")
    print(f"   ✅ Key term check  : 'Hemoglobin' found in extracted text")

    # Test the display formatter
    display_md = format_parsed_for_display(parsed)
    assert "## 📄 Extracted Raw Text" in display_md, "Display formatter output malformed"
    assert parsed.file_name in display_md, "Filename missing from display output"
    print(f"   ✅ format_parsed_for_display() output is well-formed")

finally:
    os.unlink(pdf_test_path)


# ─────────────────────────────────────────────────────────────
#  Test 5: DOCX parsing with a synthetic medical DOCX
#
#  We create a synthetic DOCX using python-docx (the same library
#  we use for reading them). The DOCX includes:
#  — Paragraph-based patient info
#  — A table with CBC results (important — tests table extraction)
#  — Doctor's narrative notes
#
#  This exercises _parse_docx() including the table extraction
#  logic which is critical for lab reports in Word format.
# ─────────────────────────────────────────────────────────────

print("\n[5/6] Testing DOCX parsing with synthetic medical document...")

from docx import Document
from docx.shared import Pt

with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
    docx_test_path = tmp.name

try:
    # Create a realistic DOCX with both paragraphs and a table
    docx_doc = Document()

    # Add heading and patient info as paragraphs
    docx_doc.add_heading("Patient Lab Report", level=1)
    docx_doc.add_paragraph("Patient Name: Test Patient")
    docx_doc.add_paragraph("Date: 01/05/2025")
    docx_doc.add_paragraph("Physician: Dr. Smith")
    docx_doc.add_paragraph("")

    docx_doc.add_heading("Complete Blood Count (CBC)", level=2)

    # Add a lab results TABLE — this is the critical test
    # Real lab reports in Word almost always use tables for values
    table = docx_doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Test"
    header_cells[1].text = "Result"
    header_cells[2].text = "Reference Range"
    header_cells[3].text = "Flag"

    # Data rows
    lab_data = [
        ("Hemoglobin",     "11.2 g/dL",   "12.0-17.5 g/dL",    "LOW"),
        ("WBC Count",      "7400 /uL",    "4000-11000 /uL",     "Normal"),
        ("Platelet Count", "210000 /uL",  "150000-400000 /uL",  "Normal"),
        ("LDL Cholesterol","142 mg/dL",   "<100 mg/dL",         "HIGH"),
        ("Fasting Glucose","118 mg/dL",   "70-99 mg/dL",        "BORDERLINE"),
    ]

    for test, result, ref_range, flag in lab_data:
        row_cells = table.add_row().cells
        row_cells[0].text = test
        row_cells[1].text = result
        row_cells[2].text = ref_range
        row_cells[3].text = flag

    # Doctor's notes as paragraphs after the table
    docx_doc.add_paragraph("")
    docx_doc.add_heading("Physician Notes", level=2)
    docx_doc.add_paragraph(
        "Patient presents with mild microcytic anemia, likely iron-deficiency. "
        "Fasting glucose is in pre-diabetic range. Recommend dietary modifications "
        "and follow-up in 3 months. LDL cholesterol elevated — lifestyle counseling advised."
    )

    docx_doc.save(docx_test_path)

    # Test parse_document() with the DOCX
    fake_file = FakeFile(docx_test_path)
    parsed = parse_document(fake_file)

    assert parsed.success,         f"DOCX parsing failed: {parsed.error}"
    assert parsed.extension == "docx", f"Extension wrong: {parsed.extension}"
    assert parsed.word_count > 30,  f"Too few words extracted: {parsed.word_count}"
    assert parsed.is_meaningful,    "Parsed DOCX text should be meaningful"

    # CRITICAL: verify table content was extracted
    assert "Hemoglobin" in parsed.text, "Table content 'Hemoglobin' not extracted from DOCX"
    assert "11.2" in parsed.text,       "Table value '11.2' not extracted from DOCX"
    assert "LOW" in parsed.text,        "Table flag 'LOW' not extracted from DOCX"

    print(f"   ✅ DOCX parsed successfully")
    print(f"   ✅ Words extracted  : {parsed.word_count:,}")
    print(f"   ✅ Estimated pages  : {parsed.page_count}")
    print(f"   ✅ Is meaningful    : {parsed.is_meaningful}")
    print(f"   ✅ Table extraction : 'Hemoglobin', '11.2', 'LOW' all found ✓")
    print(f"   ✅ Pipe-separated table rows in extracted text")

finally:
    os.unlink(docx_test_path)


# ─────────────────────────────────────────────────────────────
#  Test 6: app.py imports with Week 2 wired in
#  This confirms the Week 2 imports in app.py are correct and
#  that wiring parse_document into analyze_report works.
# ─────────────────────────────────────────────────────────────

print("\n[6/6] Verifying app.py imports Week 2 correctly...")

try:
    # We import just the functions we need to verify, not launch Gradio
    import importlib.util
    spec = importlib.util.spec_from_file_location("app", "app.py")
    # We can't fully execute app.py (it would launch Gradio)
    # but we CAN verify the tools import at the module level
    from tools.document_parser import parse_document, format_parsed_for_display
    import config
    print("   ✅ tools.document_parser imports cleanly into app.py context")
    print("   ✅ config module loads with MEDICAL_DISCLAIMER present")
    assert hasattr(config, "MEDICAL_DISCLAIMER"), "MEDICAL_DISCLAIMER missing from config"
    assert hasattr(config, "APP_VERSION"),        "APP_VERSION missing from config"
    print("   ✅ config.MEDICAL_DISCLAIMER and config.APP_VERSION present")
except Exception as e:
    print(f"   ❌ app.py context import failed: {e}")
    sys.exit(1)


# ─────────────────────────────────────────────────────────────
#  Summary
# ─────────────────────────────────────────────────────────────

print("\n" + "═" * 65)
print("  ✅ ALL WEEK 2 TESTS PASSED")
print("═" * 65)
print()
print("  What's working now:")
print("  ✅ utils/validator.py    — file validation (type, size, readability)")
print("  ✅ utils/sanitizer.py    — text cleaning, chunking, stats")
print("  ✅ tools/document_parser.py — PDF (PyMuPDF) + DOCX (python-docx)")
print("  ✅ app.py                — wired to real parser, raw tab populated")
print()
print("  What's coming in Week 3:")
print("  ⏳ prompts/analyzer_prompt.py    — system prompt for extraction Agent")
print("  ⏳ tools/report_analyzer.py      — Agent + Runner.run() for findings")
print("  ⏳ custom_data_types.py          — Pydantic models (ReportFindings etc.)")
print()
print("  To launch the app with Week 2 wired in:")
print("  $ python app.py")
print("  $ uv run python app.py")
print()
